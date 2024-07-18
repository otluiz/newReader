import os
import re
import time
from datetime import datetime
from gtts import gTTS
from pydub import AudioSegment
from docx import Document
from tkinter import Tk, Text, Button, filedialog, END, INSERT, DISABLED, NORMAL, PhotoImage, Toplevel, Label, Scrollbar, RIGHT, Y
import pygame
import PyPDF2

# Configurações do projeto
PROJECT_DIR = "project_files"
BAGS_DIR = os.path.join(PROJECT_DIR, "bags")
ICON_PATH = os.path.join(PROJECT_DIR, "play_icon.png")

# Criação dos diretórios do projeto
if not os.path.exists(PROJECT_DIR):
    os.makedirs(PROJECT_DIR)
if not os.path.exists(BAGS_DIR):
    os.makedirs(BAGS_DIR)

def replace_abbreviations(text):
    """Substitui abreviações específicas no texto"""
    text = re.sub(r'\bArt\.?\s*(\d+)', r'Artigo \1', text)
    return text

def load_bags():
    """Carrega as bags de palavras do diretório 'bags'"""
    bags = {}
    for filename in os.listdir(BAGS_DIR):
        filepath = os.path.join(BAGS_DIR, filename)
        with open(filepath, 'r', encoding='utf-8') as file:
            for line in file:
                if '=' in line:
                    key, value = line.strip().split('=', 1)
                    bags[key.strip()] = value.strip()
    return bags

def apply_bags(text, bags):
    """Aplica as substituições das bags de palavras no texto"""
    for key, value in bags.items():
        text = re.sub(r'\b' + re.escape(key) + r'\b', value, text)
    return text

def split_text(text, max_length=500):
    """Divide o texto em partes menores para evitar erros com o gTTS"""
    parts = []
    while len(text) > max_length:
        split_index = text.rfind(' ', 0, max_length)
        if split_index == -1:
            split_index = max_length
        parts.append(text[:split_index].strip())
        text = text[split_index:].strip()
    parts.append(text)
    return parts

def text_to_speech_with_highlight(text, output_filename):
    """Converte o texto em áudio, salva em um arquivo e gera um log de destaque"""
    text_parts = split_text(text)
    temp_files = []
    log_filename = os.path.join(PROJECT_DIR, f"{os.path.splitext(output_filename)[0]}_highlight.log")
    
    start_time = time.time()
    with open(log_filename, "w") as log_file:
        for i, part in enumerate(text_parts):
            tts = gTTS(part, lang='pt')
            temp_filename = os.path.join(PROJECT_DIR, f"temp_{i}.mp3")
            tts.save(temp_filename)
            temp_files.append(temp_filename)
            part_preview = part[:50].replace('\n', ' ')
            log_file.write(f"{time.time() - start_time:.2f} - {part_preview}\n")
    
    combined = AudioSegment.empty()
    for temp_file in temp_files:
        combined += AudioSegment.from_mp3(temp_file)
        os.remove(temp_file)
    
    output_filepath = os.path.join(PROJECT_DIR, output_filename)
    combined.export(output_filepath, format="mp3")
    return log_filename

def docx_to_text(docx_filename):
    """Lê um arquivo .docx e retorna seu conteúdo como uma string"""
    doc = Document(docx_filename)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def text_to_docx(docx_filename, text):
    """Salva um texto em um arquivo .docx"""
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    doc.save(docx_filename)

def pdf_to_text(pdf_filename):
    """Lê um arquivo PDF e retorna seu conteúdo como uma string"""
    pdf_text = []
    with open(pdf_filename, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            pdf_text.append(page.extract_text())
    return '\n'.join(pdf_text)

def load_file():
    """Carrega um arquivo e insere o texto na caixa de texto"""
    filepath = filedialog.askopenfilename(filetypes=[("Document files", "*.docx *.txt *.pdf")])
    if filepath:
        ext = os.path.splitext(filepath)[1].lower()
        if ext == ".docx":
            text = docx_to_text(filepath)
        elif ext == ".pdf":
            text = pdf_to_text(filepath)
        elif ext == ".txt":
            with open(filepath, 'r', encoding='utf-8') as file:
                text = file.read()
        else:
            status_label.config(text="Formato de arquivo não suportado!")
            return

        text = replace_abbreviations(text)
        bags = load_bags()
        text = apply_bags(text, bags)
        text_box.delete(1.0, END)
        text_box.insert(INSERT, text)
        generate_button.config(state=NORMAL)
        play_button.config(state=DISABLED)
        save_button.config(state=NORMAL)
        status_label.config(text=f"Carregado: {os.path.basename(filepath)}")
        root.filepath = filepath  # Salva o caminho do arquivo carregado

def play_audio_with_highlight(log_filename, audio_filename):
    """Reproduz o áudio e destaca o texto conforme o log"""
    with open(log_filename, "r") as log_file:
        lines = log_file.readlines()
    
    highlight_intervals = []
    highlight_texts = []
    for line in lines:
        if " - " in line:
            timestamp, text_preview = line.split(" - ", 1)
            try:
                highlight_intervals.append(float(timestamp))
                highlight_texts.append(text_preview.strip())
            except ValueError:
                continue
    
    pygame.mixer.init()
    pygame.mixer.music.load(audio_filename)
    pygame.mixer.music.play()
    
    start_time = time.time()
    for interval, text_preview in zip(highlight_intervals, highlight_texts):
        while time.time() - start_time < interval:
            time.sleep(0.1)
        start_idx = text_box.search(text_preview, 1.0, stopindex=END)
        if start_idx:
            end_idx = f"{start_idx}+{len(text_preview)}c"
            text_box.tag_remove("highlight", 1.0, END)
            text_box.tag_add("highlight", start_idx, end_idx)
            text_box.tag_config("highlight", background="yellow")
            text_box.see(start_idx)

def stop_audio():
    """Para a reprodução do áudio"""
    pygame.mixer.music.stop()
    text_box.tag_remove("highlight", 1.0, END)

def save_file():
    """Salva o texto editado de volta no arquivo original"""
    text = text_box.get(1.0, END)
    if hasattr(root, 'filepath'):
        ext = os.path.splitext(root.filepath)[1].lower()
        if ext == ".docx":
            text_to_docx(root.filepath, text)
        elif ext == ".txt":
            with open(root.filepath, 'w', encoding='utf-8') as file:
                file.write(text)
        status_label.config(text=f"Salvo: {os.path.basename(root.filepath)}")

def generate_audio():
    """Gera o áudio e o log de destaque"""
    text = text_box.get(1.0, END)
    if hasattr(root, 'filepath'):
        base_filename = os.path.splitext(os.path.basename(root.filepath))[0]
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        output_filename = f"{base_filename}_{timestamp}.mp3"
    else:
        output_filename = "output.mp3"
    log_filename = os.path.join(PROJECT_DIR, f"{os.path.splitext(output_filename)[0]}_highlight.log")
    
    status_label.config(text="Audio está sendo gerado, aguarde...")
    root.update_idletasks()  # Atualiza a interface para mostrar a mensagem

    if not (os.path.exists(os.path.join(PROJECT_DIR, output_filename)) and os.path.exists(log_filename)):
        log_filename = text_to_speech_with_highlight(text, output_filename)
    
    status_label.config(text="Audio gerado com sucesso!")
    play_button.config(state=NORMAL)
    root.output_filename = output_filename  # Salva o nome do arquivo de saída

def play_audio():
    """Inicia a reprodução do áudio com destaque"""
    if hasattr(root, 'output_filename'):
        output_filename = os.path.join(PROJECT_DIR, root.output_filename)
        log_filename = os.path.join(PROJECT_DIR, f"{os.path.splitext(root.output_filename)[0]}_highlight.log")
        play_audio_with_highlight(log_filename, output_filename)

# Interface Gráfica com Tkinter
root = Tk()
root.title("Text to Speech with Highlight")

scrollbar = Scrollbar(root)
scrollbar.pack(side=RIGHT, fill=Y)

text_box = Text(root, wrap='word', width=100, height=30, yscrollcommand=scrollbar.set)
text_box.pack(pady=20)

scrollbar.config(command=text_box.yview)

button_frame = Button(root)
button_frame.pack(pady=10)

load_button = Button(button_frame, text="Load File", command=load_file)
load_button.pack(side='left', padx=10)

generate_button = Button(button_frame, text="Generate Audio", command=generate_audio, state=DISABLED)
generate_button.pack(side='left', padx=10)

play_button = Button(button_frame, text="Play", command=play_audio, state=DISABLED)
play_button.pack(side='left', padx=10)

stop_button = Button(button_frame, text="Stop", command=stop_audio)
stop_button.pack(side='left', padx=10)

save_button = Button(button_frame, text="Save", command=save_file, state=DISABLED)
save_button.pack(side='left', padx=10)

# Adiciona o rótulo de status
status_label = Label(root, text="")
status_label.pack(pady=10)

# Função para criar um tooltip
def create_tooltip(widget, text):
    tooltip = Toplevel(widget)
    tooltip.wm_overrideredirect(True)
    tooltip.wm_geometry("+0+0")
    label = Label(tooltip, text=text, background="yellow", relief="solid", borderwidth=1)
    label.pack()

play_icon = None
try:
    play_icon = PhotoImage(file=ICON_PATH)
except Exception as e:
    print(f"Erro ao carregar ícone: {e}")

root.mainloop()
